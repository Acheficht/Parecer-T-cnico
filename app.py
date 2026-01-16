import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from fpdf import FPDF
import io
import json
import os
import tempfile
import base64
from datetime import date
from PIL import Image

# --- 1. CONFIGURAÇÃO INICIAL E CSS ---
st.set_page_config(page_title="Parecer Técnico v12 (Fixed)", layout="wide")

st.markdown("""
<style>
    div[data-testid="stVerticalBlock"] > div > button {
        border-color: #ff4b4b;
        color: #ff4b4b;
    }
    div[data-testid="stVerticalBlock"] > div > button:hover {
        background-color: #ff4b4b;
        color: white;
    }
    .img-container {
        border: 1px solid #ddd;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 10px;
        text-align: center;
        background-color: #f9f9f9;
    }
    .block-container { padding-top: 2rem; }
</style>
""", unsafe_allow_html=True)

# --- 2. LISTA MESTRA DE OPÇÕES ---
OPCOES_LISTA = [
    "Inconsistências em Ficha do imóvel", "Inconsistências em Sobreposição com outros IRs",
    "Outras Sobreposições", "Inconsistências em Áreas embargadas", "Inconsistências em Assentamentos",
    "Inconsistências em UC", "Inconsistências em Cobertura do solo", 
    "Inconsistências em Infraestrutura e utilidade pública",
    "Inconsistências em Reservatório para abastecimento ou geração de energia",
    "Inconsistências em APP hidrografia", "Inconsistências em APP Relevo", "Inconsistências em Uso restrito",
    "Inconsistências em outras APPs", "Inconsistências em RL averbada, RL aprovada e não averbada",
    "Inconsistências em Área de RL exigida por lei", "Inconsistências em Localização e cobertura do solo",
    "Inconsistências em Regularidade do IR", "Observação"
]

# --- 3. GERENCIAMENTO DE ESTADO ---
if 'dados' not in st.session_state:
    st.session_state['dados'] = {
        "car": "", "sp_not": "", "imovel": "", 
        "nome": "", "doc": "", "cidade": "Mogi das Cruzes",
        "selecionados": [], "textos": {},
        "imagens_b64": {} 
    }

if 'uploader_ids' not in st.session_state:
    st.session_state['uploader_ids'] = {}

# --- 4. FUNÇÕES DE CALLBACK (A MÁGICA ACONTECE AQUI) ---
def processar_upload(item):
    """
    Esta função é chamada AUTOMATICAMENTE assim que um arquivo é escolhido.
    Ela processa o arquivo e muda a chave do uploader para limpá-lo.
    """
    # 1. Descobre qual é a chave atual deste uploader
    uid = st.session_state['uploader_ids'].get(item, 0)
    key_widget = f"uploader_{item}_{uid}"
    
    # 2. Pega o arquivo direto do session_state
    uploaded_file = st.session_state.get(key_widget)
    
    if uploaded_file:
        try:
            # Processa e Salva
            uploaded_file.seek(0)
            bytes_data = uploaded_file.read()
            b64_str = base64.b64encode(bytes_data).decode('utf-8')
            
            if item not in st.session_state['dados']['imagens_b64']:
                st.session_state['dados']['imagens_b64'][item] = []
            
            st.session_state['dados']['imagens_b64'][item].append(b64_str)
            
            # 3. INCREMENTA O ID
            # Isso faz com que, ao recarregar a página, o Streamlit crie um NOVO uploader vazio
            st.session_state['uploader_ids'][item] = uid + 1
            
        except Exception as e:
            st.error(f"Erro no processamento: {e}")

# --- 5. OUTRAS FUNÇÕES ---

def obter_data_extenso():
    meses = {1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril', 5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto', 9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'}
    hj = date.today()
    return f"{hj.day} de {meses[hj.month]} de {hj.year}"

def limpar_tudo():
    st.session_state['dados'] = {
        "car": "", "sp_not": "", "imovel": "", 
        "nome": "", "doc": "", "cidade": "", 
        "selecionados": [], "textos": {},
        "imagens_b64": {}
    }
    st.session_state['uploader_ids'] = {}
    
    campos = ["car", "sp_not", "imovel", "nome", "doc", "cidade"]
    for c in campos:
        st.session_state[f"input_{c}"] = ""
        
    for op in OPCOES_LISTA:
        st.session_state[f"chk_{op}"] = False

def limpar_campo_cabecalho(key_sulfix):
    st.session_state[f"input_{key_sulfix}"] = ""

def limpar_conteudo_item(item):
    """
    CORREÇÃO REALIZADA: Limpa apenas o texto, mantém as imagens.
    """
    # Remove do dicionário de persistência
    if item in st.session_state['dados']['textos']:
        del st.session_state['dados']['textos'][item]
    
    # Limpa visualmente o widget de texto
    st.session_state[f"txt_area_{item}"] = ""
    
    # OBS: A parte que deletava 'imagens_b64' foi removida para atender ao pedido.

def remover_imagem_especifica(item, index):
    if item in st.session_state['dados']['imagens_b64']:
        lista = st.session_state['dados']['imagens_b64'][item]
        if 0 <= index < len(lista):
            lista.pop(index)
            if not lista:
                del st.session_state['dados']['imagens_b64'][item]

def toggle_item(item_nome):
    chave = f"chk_{item_nome}"
    estado = st.session_state.get(chave, False)
    lista = st.session_state['dados']['selecionados']
    
    if estado and item_nome not in lista:
        lista.append(item_nome)
    elif not estado and item_nome in lista:
        lista.remove(item_nome)
    st.session_state['dados']['selecionados'] = lista

def formatar_documento(n):
    if not n: return ""
    n = "".join(filter(str.isdigit, str(n)))
    if len(n) == 11: return f"{n[:3]}.{n[3:6]}.{n[6:9]}-{n[9:]}"
    elif len(n) == 14: return f"{n[:2]}.{n[2:5]}.{n[5:8]}/{n[8:12]}-{n[12:]}"
    return n

def b64_para_tempfile(b64_str):
    try:
        bytes_data = base64.b64decode(b64_str)
        file_obj = io.BytesIO(bytes_data)
        img = Image.open(file_obj)
        if img.mode in ('RGBA', 'LA'):
            background = Image.new(img.mode[:-1], img.size, (255, 255, 255))
            background.paste(img, img.split()[-1])
            img = background
        img = img.convert('RGB')
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tf:
            img.save(tf, format='JPEG', quality=95)
            return tf.name
    except Exception:
        return None

# --- 6. GERAÇÃO DE PDF ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', '', 10)
        self.set_xy(-20, 20) 
        self.cell(0, 0, str(self.page_no()), 0, 0, 'R')
        self.ln(20)

def gerar_pdf_bytes():
    try:
        pdf = PDF()
        pdf.set_margins(30, 30, 20)
        pdf.add_page()
        def safe_text(text): return text.encode('latin-1', 'replace').decode('latin-1') if text else ""

        pdf.set_font("Arial", 'B', 14); pdf.cell(0, 8, safe_text("Justificativa do Parecer Técnico"), ln=True, align='C')
        pdf.set_font("Arial", 'B', 12); pdf.cell(0, 6, safe_text(f"CAR: {st.session_state['dados']['car']}"), ln=True, align='C')
        pdf.cell(0, 6, safe_text(f"SP-NOT: {st.session_state['dados']['sp_not']}"), ln=True, align='C'); pdf.ln(5)
        
        pdf.set_font("Arial", 'B', 12); pdf.write(6, safe_text("Nome do Imóvel Rural: ")); pdf.set_font("Arial", '', 12); pdf.write(6, safe_text(st.session_state['dados']['imovel'])); pdf.ln(6)
        pdf.set_font("Arial", 'B', 12); pdf.write(6, safe_text("Nome: ")); pdf.set_font("Arial", '', 12); pdf.write(6, safe_text(st.session_state['dados']['nome'])); pdf.ln(6)
        pdf.set_font("Arial", 'B', 12); pdf.write(6, safe_text("CPF/CNPJ: ")); pdf.set_font("Arial", '', 12); pdf.write(6, safe_text(formatar_documento(st.session_state['dados']['doc']))); pdf.ln(10)
        
        for i, item in enumerate(st.session_state['dados']['selecionados']):
            pdf.set_font("Arial", 'B', 12); pdf.cell(0, 8, safe_text(f"{i+1}. {item}"), ln=True)
            pdf.set_font("Arial", '', 12)
            
            # Recupera texto (prioriza widget, fallback para dados salvos)
            texto_raw = st.session_state.get(f"txt_area_{item}", st.session_state['dados']['textos'].get(item, ""))
            pdf.multi_cell(0, 7, "      " + safe_text(texto_raw), align='J')
            
            lista_imgs = st.session_state['dados']['imagens_b64'].get(item, [])
            if lista_imgs:
                pdf.ln(2)
                for b64_img in lista_imgs:
                    temp_path = b64_para_tempfile(b64_img)
                    if temp_path:
                        try:
                            x_pos = (210 - 120) / 2 
                            if pdf.get_y() > 220: pdf.add_page()
                            pdf.image(temp_path, x=x_pos, w=120)
                            pdf.ln(5)
                        finally:
                            if os.path.exists(temp_path): os.unlink(temp_path)
                pdf.ln(2)
            pdf.ln(4)
            
        pdf.ln(10); pdf.cell(0, 6, "________________________________________", ln=True, align='R')
        pdf.cell(0, 6, safe_text(st.session_state['dados']['nome']), ln=True, align='R')
        cidade_doc = st.session_state['dados']['cidade'] if st.session_state['dados']['cidade'] else "Mogi das Cruzes"
        pdf.cell(0, 6, safe_text(f"{cidade_doc}, {obter_data_extenso()}."), ln=True, align='R')
        return pdf.output(dest='S').encode('latin-1', 'replace')
    except Exception as e: return None

# --- 7. BARRA LATERAL ---
with st.sidebar:
    st.header("🗂️ Arquivos")
    arquivo_upload = st.file_uploader("📂 Carregar Trabalho (.json)", type=["json"])
    
    if arquivo_upload is not None:
        if st.button("🔄 Confirmar Carregamento"):
            try:
                dados_carregados = json.load(arquivo_upload)
                
                imagens = dados_carregados.get("imagens_b64", {})
                for k, v in imagens.items():
                    if isinstance(v, str): imagens[k] = [v]
                dados_carregados["imagens_b64"] = imagens

                st.session_state['dados'] = dados_carregados
                
                for k in ["car", "sp_not", "imovel", "nome", "doc", "cidade"]:
                    st.session_state[f"input_{k}"] = dados_carregados.get(k, "")
                
                selecionados = dados_carregados.get("selecionados", [])
                for op in OPCOES_LISTA:
                    st.session_state[f"chk_{op}"] = (op in selecionados)
                
                # Carrega textos para a session_state de widgets também, por garantia
                for item, texto in dados_carregados.get("textos", {}).items():
                    st.session_state[f"txt_area_{item}"] = texto
                
                st.session_state['uploader_ids'] = {} 
                st.success("✅ Carregado com Sucesso!")
                st.rerun()
            except Exception as e:
                st.error(f"Erro: {e}")
    
    st.markdown("---")
    dados_download = json.dumps(st.session_state['dados'], indent=4)
    st.download_button("💾 Salvar Backup", dados_download, "backup_multi_imagens.json", "application/json")
    st.markdown("---")
    st.button("🗑️ Limpar Tudo", on_click=limpar_tudo, type="primary")

# --- 8. INTERFACE PRINCIPAL ---
st.title("📄 Gerador de Parecer Técnico (Multi Imagens)")

tab_edit, tab_preview = st.tabs(["✍️ Edição", "👁️ Pré-visualização Real (PDF)"])

with tab_edit:
    st.subheader("1. Cabeçalho")
    def campo_com_lixeira(label, key_suffix):
        c_input, c_btn = st.columns([0.85, 0.15])
        with c_input:
            val = st.text_input(label, key=f"input_{key_suffix}")
            st.session_state['dados'][key_suffix] = val
        with c_btn:
            st.write(""); st.write("")
            st.button("🗑️", key=f"del_header_{key_suffix}", on_click=limpar_campo_cabecalho, args=(key_suffix,))

    c1, c2 = st.columns(2)
    with c1:
        campo_com_lixeira("CAR:", "car")
        campo_com_lixeira("SP-NOT (Número):", "sp_not")
        campo_com_lixeira("Nome do Imóvel:", "imovel")
    with c2:
        campo_com_lixeira("Nome do Requerente:", "nome")
        campo_com_lixeira("CPF/CNPJ:", "doc")
        campo_com_lixeira("Cidade:", "cidade")

    st.markdown("---")
    st.subheader("2. Seleção de Inconsistências")
    
    cols_check = st.columns(3)
    for i, op in enumerate(OPCOES_LISTA):
        cols_check[i%3].checkbox(op, key=f"chk_{op}", on_change=toggle_item, args=(op,))

    selecionados = st.session_state['dados']['selecionados']
    
    if selecionados:
        st.markdown("---")
        st.info("Preencha os detalhes abaixo:")
        
        for item in selecionados:
            col_expander, col_lixo = st.columns([0.92, 0.08])
            
            with col_lixo:
                st.write("")
                # Botão que agora limpa SÓ o texto
                st.button("🗑️", key=f"clean_content_{item}", on_click=limpar_conteudo_item, args=(item,), help="Limpar apenas o texto")

            with col_expander:
                with st.expander(f"📝 {item}", expanded=True):
                    chave_txt = f"txt_area_{item}"
                    
                    # --- CORREÇÃO DE PERSISTÊNCIA DE TEXTO AQUI ---
                    # Tentamos pegar o valor persistente do dicionário 'dados'.
                    # Se não existir, aí sim tentamos ver se o widget tem algo (o que é raro se foi recém montado)
                    val_inicial = st.session_state['dados']['textos'].get(item, "")
                    
                    texto = st.text_area("Descrição:", value=val_inicial, key=chave_txt, height=100)
                    
                    # Atualiza o dicionário principal sempre que digita
                    if texto: st.session_state['dados']['textos'][item] = texto
                    # Se o usuário apagar tudo manualmente, removemos do dict
                    elif item in st.session_state['dados']['textos'] and not texto:
                         del st.session_state['dados']['textos'][item]
                    
                    st.markdown("#### 🖼️ Imagens do Item")
                    
                    # 1. GALERIA
                    lista_imgs = st.session_state['dados']['imagens_b64'].get(item, [])
                    if lista_imgs:
                        cols_imgs = st.columns(3)
                        for idx, img_b64 in enumerate(lista_imgs):
                            with cols_imgs[idx % 3]:
                                st.markdown("<div class='img-container'>", unsafe_allow_html=True)
                                st.image(base64.b64decode(img_b64), use_container_width=True)
                                if st.button("❌", key=f"del_img_{item}_{idx}"):
                                    remover_imagem_especifica(item, idx)
                                    st.rerun()
                                st.markdown("</div>", unsafe_allow_html=True)
                    else:
                        st.caption("Nenhuma imagem adicionada ainda.")

                    # 2. UPLOAD (COM CALLBACK)
                    uid = st.session_state['uploader_ids'].get(item, 0)
                    key_uploader = f"uploader_{item}_{uid}"
                    
                    st.file_uploader(
                        f"Adicionar nova imagem em '{item}'", 
                        type=['png','jpg','jpeg'], 
                        key=key_uploader,
                        on_change=processar_upload, 
                        args=(item,)
                    )

with tab_preview:
    st.info("Visualização exata do PDF final.")
    pdf_bytes_preview = gerar_pdf_bytes()
    if pdf_bytes_preview:
        base64_pdf = base64.b64encode(pdf_bytes_preview).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
    else:
        st.warning("Preencha os dados para gerar a visualização.")

# --- DOWNLOADS ---
st.markdown("---")
st.subheader("🚀 Baixar Arquivos")
col_d1, col_d2 = st.columns(2)
nome_safe = st.session_state['dados']['nome'].strip() or "Parecer"
nome_arquivo = "".join([c for c in nome_safe if c.isalnum() or c in (' ','-','_')]).strip()

if st.session_state['dados']['car']:
    # WORD
    doc = Document()
    sec = doc.sections[0]; sec.top_margin = Cm(3); sec.bottom_margin = Cm(2); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(12)
    
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Justificativa do Parecer Técnico"); r.bold=True; r.font.size=Pt(14)
    p = doc.add_paragraph(f"CAR: {st.session_state['dados']['car']}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True
    p = doc.add_paragraph(f"SP-NOT: {st.session_state['dados']['sp_not']}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True
    doc.add_paragraph()
    
    p = doc.add_paragraph(); p.add_run("Nome do Imóvel Rural: ").bold=True; p.add_run(st.session_state['dados']['imovel'])
    p = doc.add_paragraph(); p.add_run("Nome: ").bold=True; p.add_run(st.session_state['dados']['nome'])
    p = doc.add_paragraph(); p.add_run("CPF/CNPJ: ").bold=True; p.add_run(formatar_documento(st.session_state['dados']['doc']))
    doc.add_paragraph()
    
    for i, item in enumerate(st.session_state['dados']['selecionados']):
        p = doc.add_paragraph(f"{i+1}. {item}"); p.runs[0].bold=True
        
        # Recupera texto (prioriza widget, fallback para dados salvos)
        texto_item = st.session_state.get(f"txt_area_{item}", st.session_state['dados']['textos'].get(item, ""))
        
        p = doc.add_paragraph(texto_item); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent = Cm(1.25)
        
        lista_imgs = st.session_state['dados']['imagens_b64'].get(item, [])
        for b64_img in lista_imgs:
            temp_path = b64_para_tempfile(b64_img)
            if temp_path:
                try:
                    doc.add_paragraph()
                    doc.add_picture(temp_path, width=Cm(14))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception: pass
                finally:
                    if os.path.exists(temp_path): os.unlink(temp_path)

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph("________________________________________"); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(st.session_state['dados']['nome']); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cidade_doc = st.session_state['dados']['cidade'] if st.session_state['dados']['cidade'] else "Mogi das Cruzes"
    p = doc.add_paragraph(f"{cidade_doc}, {obter_data_extenso()}."); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    buffer_word = io.BytesIO(); doc.save(buffer_word); buffer_word.seek(0)
    with col_d1:
        st.download_button("⬇️ Baixar Word (.docx)", buffer_word, f"{nome_arquivo}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if pdf_bytes_preview:
    with col_d2:
        st.download_button("⬇️ Baixar PDF (.pdf)", pdf_bytes_preview, f"{nome_arquivo}.pdf", "application/pdf")